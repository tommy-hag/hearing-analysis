#!/usr/bin/env python3
"""
DOCX builder used by the webapp.
- Inserts a standard template block under each H2 section (idempotent via marker)
- Parses Markdown with CriticMarkup highlights/comments and converts to Word
- Preserves bold (**text**) and italic (*text*)
- Rebuilds the TOC in the provided Word template
- Enables update-on-open so TOC refreshes when the document is opened in Word

CLI usage (as invoked by the server):

  python3 build_docx_colab.py \
      --markdown - \
      --out /abs/path/output.docx \
      --template /abs/path/Bilag 6 Svar på henvendelser i høringsperioden.docx \
      --template-block /abs/path/blok.md

If --markdown is '-' content is read from STDIN (UTF-8).
"""

import argparse
import os
import sys
import re
import hashlib
from pathlib import Path

def _bootstrap_paths():
    # Ensure local package directories from PYTHONPATH are on sys.path (handles colon-separated lists)
    try:
        pp = os.environ.get('PYTHONPATH', '')
        if pp:
            for p in pp.split(os.pathsep):
                p = p.strip()
                if p and p not in sys.path:
                    sys.path.insert(0, p)
    except Exception:
        pass
    # Also try common project-local python_packages folders
    try:
        here = Path(__file__).resolve()
        candidates = [
            here.parent / 'vendor',
            here.parent.parent / 'python_packages',               # fetcher/python_packages
            here.parent.parent.parent / 'python_packages'          # repo-root/python_packages (if any)
        ]
        for c in candidates:
            try:
                if c and c.exists():
                    p = str(c)
                    if p not in sys.path:
                        sys.path.insert(0, p)
            except Exception:
                pass
    except Exception:
        pass


def _try_runtime_install():
    try:
        import subprocess
        here = Path(__file__).resolve().parent
        target = here / 'vendor'
        try:
            target.mkdir(parents=True, exist_ok=True)
        except Exception:
            pass
        cmd = [sys.executable, '-m', 'pip', 'install', '--no-cache-dir', '--no-warn-script-location', '--prefer-binary', '--only-binary', ':all:', '--upgrade', '--target', str(target), 'python-docx>=1.2.0', 'lxml>=5', 'Pillow>=10.0.0']
        subprocess.run(cmd, check=False, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        if str(target) not in sys.path:
            sys.path.insert(0, str(target))
    except Exception:
        pass


def _import_docx_or_die():
    global python_docx_module, Document, OxmlElement, qn, Paragraph
    try:
        import docx as python_docx_module  # python-docx >= 1.2.0 recommended
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.text.paragraph import Paragraph
        return True
    except Exception as e:
        # One-shot runtime install fallback, then retry
        _try_runtime_install()
        try:
            import docx as python_docx_module
            from docx import Document
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.text.paragraph import Paragraph
            return True
        except Exception as e2:
            print(f"FATAL: Mangler python-docx ({e2}). Installér med: pip install 'python-docx>=1.2.0'", file=sys.stderr)
            return False


_bootstrap_paths()
if not _import_docx_or_die():
    sys.exit(2)


# --- Helpers ---

def _norm(s: str) -> str:
    return s.replace("\r\n", "\n").replace("\r", "\n")


def strip_html_comments(md_text: str) -> str:
    return re.sub(r'<!--.*?-->', '', md_text, flags=re.DOTALL)


def preprocess_and_extract_criticmarkup(md_text: str):
    critic_blocks = []

    def replacement_func(match):
        highlight_content, comment_content = match.group(1), match.group(2)
        block_index = len(critic_blocks)
        critic_blocks.append({'highlight': highlight_content, 'comment': comment_content})
        return f"__CRITIC_MARKUP_{block_index}__"

    pattern = re.compile(r'\{==(.*?)==\}\s*\{>>(.+?)<<\}', re.DOTALL)
    return pattern.sub(replacement_func, md_text), critic_blocks


def _format_text_in_runs(p: Paragraph, text: str):
    token_pattern = re.compile(r'(\*\*([^*]+?)\*\*)|(\*([^*]+?)\*)')
    last_end = 0
    for match in token_pattern.finditer(text):
        p.add_run(text[last_end:match.start()])
        if match.group(2):
            p.add_run(match.group(2)).bold = True
        elif match.group(4):
            p.add_run(match.group(4)).italic = True
        last_end = match.end()
    p.add_run(text[last_end:])


def _normalize_comment_text(text: str) -> str:
    # Normalize Windows/Mac line endings and convert literal \n to newlines
    s = text.replace('\r\n', '\n').replace('\r', '\n')
    s = s.replace('\\n', '\n')
    return s


def _try_apply_style(p: Paragraph, doc: Document, style_names: list[str]):
    for name in style_names:
        try:
            s = doc.styles[name]
            if s is not None:
                p.style = s
                return True
        except Exception:
            continue
    return False


def _add_comment_safe(doc: Document, runs, comment_text: str):
    """Attach a Word comment to the given runs if supported; fallback to inline note."""
    comment_text_norm = _normalize_comment_text(comment_text)
    try:
        add_comment = getattr(doc, 'add_comment', None)
        if callable(add_comment) and runs:
            c = add_comment(text='', author="Bliv Hørt AI", initials="AI", runs=runs)
            c.paragraphs[0].clear()
            for i, line in enumerate(comment_text_norm.split('\n')):
                if i > 0:
                    c.paragraphs[0].add_run().add_break()
                _format_text_in_runs(c.paragraphs[0], line)
            return True
    except Exception:
        pass
    try:
        # Fallback: append inline note to the same paragraph to avoid extra line breaks
        if runs:
            p = runs[-1].paragraph
            lines = comment_text_norm.split('\n')
            if lines:
                _format_text_in_runs(p, f" [Kommentar] {lines[0]}")
                for extra in lines[1:]:
                    p.add_run().add_break()
                    _format_text_in_runs(p, extra)
            return True
    except Exception:
        pass
    return False


def process_line_content(doc: Document, p: Paragraph, line_text: str, critic_blocks: list):
    placeholder_pattern = re.compile(r'__CRITIC_MARKUP_(\d+)__')
    last_end = 0
    for match in placeholder_pattern.finditer(line_text):
        plain_text_before = line_text[last_end:match.start()]
        if plain_text_before:
            _format_text_in_runs(p, plain_text_before)

        block_index = int(match.group(1))
        if 0 <= block_index < len(critic_blocks):
            block = critic_blocks[block_index]
            runs_before_count = len(p.runs)
            _format_text_in_runs(p, block['highlight'])
            runs_to_comment_on = p.runs[runs_before_count:]

            processed_comment_text = re.sub(r'(\*\*Henvendelse \d+\*\*)\s*(")', r'\1\n\2', block['comment'])
            _add_comment_safe(doc, runs_to_comment_on, processed_comment_text)

        last_end = match.end()

    plain_text_after = line_text[last_end:]
    if plain_text_after:
        _format_text_in_runs(p, plain_text_after)


# --- TOC and structure ---

def enable_update_fields_on_open(doc: Document):
    settings = doc.settings.element
    for el in settings.findall(qn('w:updateFields')):
        settings.remove(el)
    update_fields = OxmlElement('w:updateFields')
    update_fields.set(qn('w:val'), 'true')
    settings.append(update_fields)


def set_paragraph_outline_level(paragraph: Paragraph, level: int):
    pPr = paragraph._p.get_or_add_pPr()
    for el in pPr.findall(qn('w:outlineLvl')):
        pPr.remove(el)
    outline = OxmlElement('w:outlineLvl')
    outline.set(qn('w:val'), str(level))
    pPr.append(outline)


def exclude_main_title_from_toc(doc: Document):
    for p in doc.paragraphs:
        if re.match(r'^\s*forslag til', p.text, flags=re.IGNORECASE):
            set_paragraph_outline_level(p, 9)
            return


def read_existing_toc_instr(doc: Document) -> str:
    body = doc._element.body
    fld = body.xpath('.//w:fldSimple[contains(@w:instr, "TOC")]')
    if fld:
        return fld[0].get(qn('w:instr'))
    instr_nodes = body.xpath('.//w:instrText[contains(text(), "TOC")]')
    if instr_nodes:
        full_instr = "".join([node.text for node in instr_nodes])
        m = re.search(r'(TOC\s.*)', full_instr.strip(), re.IGNORECASE)
        if m:
            return m.group(1).strip()
    return 'TOC \\o "1-3" \\h \\z \\u'


def find_first_toc_entry_index(doc: Document):
    for i, p in enumerate(doc.paragraphs):
        if (p.style.name or "").strip().upper().startswith('TOC'):
            return i
    for i, p in enumerate(doc.paragraphs):
        if p._p.xpath('.//w:fldSimple[contains(@w:instr, "TOC")]'):
            return i
    return None


def _pstyle_val(p_el):
    vals = p_el.xpath('./w:pPr/w:pStyle/@w:val')
    return vals[0] if vals else ''


def clear_and_replace_toc(doc: Document):
    toc_instr = read_existing_toc_instr(doc)
    start_index = find_first_toc_entry_index(doc)
    if start_index is None:
        return

    all_paras_xml = doc._element.body.xpath('./w:p')
    paras_to_delete_xml = []
    for i in range(start_index, len(all_paras_xml)):
        p_xml = all_paras_xml[i]
        style_val = (_pstyle_val(p_xml) or '').upper()
        is_toc_field = p_xml.xpath('.//w:fldSimple | .//w:fldChar')
        is_toc_style = style_val.startswith('TOC')
        if is_toc_field or is_toc_style:
            paras_to_delete_xml.append(p_xml)
        else:
            break

    if not paras_to_delete_xml:
        return

    anchor_xml = paras_to_delete_xml[0]
    toc_p_xml = OxmlElement('w:p')
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), toc_instr)
    run = OxmlElement('w:r')
    text = OxmlElement('w:t')
    text.text = 'Højreklik -> Opdater felt'
    run.append(text)
    fld.append(run)
    toc_p_xml.append(fld)
    anchor_xml.addprevious(toc_p_xml)

    for p_xml in paras_to_delete_xml:
        p_xml.getparent().remove(p_xml)


# --- Template block insertion under each H2 ---

def _sha1(text: str) -> str:
    return hashlib.sha1(text.encode('utf-8')).hexdigest()


def build_insertion_block(template_md: str):
    tmpl = _norm(template_md).strip()
    sha = _sha1(tmpl)
    marker = f"<!-- INSERTED_BLOCK_SHA1:{sha} -->"
    # No surrounding newlines; insertion happens with precise newlines in caller
    insertion = f"{tmpl} {marker}"
    return insertion, marker


def insert_block_in_sections(md_text: str, insertion: str, marker: str):
    text = _norm(md_text)
    h2_matches = list(re.finditer(r"^##\s+.*$", text, flags=re.MULTILINE))
    if not h2_matches:
        return text, 0, 0

    next_h1_or_h2 = re.compile(r"^#{1,2}\s+", flags=re.MULTILINE)
    parts = []
    cursor = 0
    inserted_count = 0

    for m in h2_matches:
        parts.append(text[cursor:m.start()])
        next_match = next_h1_or_h2.search(text, pos=m.end())
        end_pos = next_match.start() if next_match else len(text)
        section = text[m.start():end_pos]
        already_has_block = (marker in section) or ("### Forvaltningens svar" in section)
        if not already_has_block:
            # Remove any trailing newlines from the section, then ensure exactly one newline
            # before the inserted block, and no extra blank line after.
            section_no_trailing_nl = section.rstrip('\n')
            parts.append(section_no_trailing_nl)
            parts.append('\n')  # start block on a new line
            parts.append(insertion)
            # Only add a separating newline after the block if there is a next heading
            if next_match is not None:
                parts.append('\n')
            inserted_count += 1
        else:
            parts.append(section)
        cursor = end_pos

    parts.append(text[cursor:])
    return "".join(parts), inserted_count, len(h2_matches)


# --- Build ---

def build_doc_from_markdown(markdown_text: str, template_file: Path, output_file: Path,
                            template_block_file: Path | None = None):
    if not template_file.exists():
        Document().save(str(template_file))

    # Insert template blocks (if provided)
    if template_block_file and template_block_file.exists():
        tmpl_md = template_block_file.read_text(encoding='utf-8')
        insertion, marker = build_insertion_block(tmpl_md)
        markdown_text, _inserted, _h2 = insert_block_in_sections(markdown_text, insertion, marker)

    # CriticMarkup pre-processing
    md_text = strip_html_comments(markdown_text)
    sanitized_text, critic_blocks = preprocess_and_extract_criticmarkup(md_text)

    # Build document content
    doc = Document(str(template_file))
    exclude_main_title_from_toc(doc)
    clear_and_replace_toc(doc)

    def _apply_heading_style_by_level(doc: Document, p: Paragraph, level: int):
        # Prefer explicit localized heading styles rather than relying on python-docx defaults
        # This improves consistency across localized templates (e.g., Danish 'Overskrift {n}')
        if level < 1:
            level = 1
        candidate_styles = [
            f'Heading {level}',
            f'Overskrift {level}',
            f'Rubrik {level}',
        ]
        # Only for level 1, also try generic Title/Rubrik/Titel if present in template
        if level == 1:
            candidate_styles.extend(['Title', 'Titel', 'Rubrik'])
        _try_apply_style(p, doc, candidate_styles)

    heading_pattern = re.compile(r'^\s{0,3}(#{1,6})\s*(.*?)\s*#*\s*$')

    for raw_line in sanitized_text.splitlines():
        line = raw_line.rstrip('\n')
        if not line.strip():
            p = doc.add_paragraph()
            _try_apply_style(p, doc, ['Normal', 'Brødtekst', 'Body Text'])
            continue
        m = heading_pattern.match(line)
        if m:
            level = len(m.group(1))
            # Use a plain paragraph and apply style explicitly for better locale compatibility
            p = doc.add_paragraph()
            set_paragraph_outline_level(p, max(0, level - 1))
            process_line_content(doc, p, m.group(2).strip(), critic_blocks)
            _apply_heading_style_by_level(doc, p, level)
        else:
            p = doc.add_paragraph()
            process_line_content(doc, p, line.strip(), critic_blocks)
            _try_apply_style(p, doc, ['Normal', 'Brødtekst', 'Body Text'])

    enable_update_fields_on_open(doc)
    doc.save(str(output_file))


def main(argv=None):
    ap = argparse.ArgumentParser(description='Colab-aligned DOCX builder')
    ap.add_argument('--markdown', required=True, help='"-" to read from stdin, or a file path')
    ap.add_argument('--out', required=True, help='Output .docx path')
    ap.add_argument('--template', required=True, help='Template .docx path')
    ap.add_argument('--template-block', required=True, help='Template block .md path')
    args = ap.parse_args(argv)

    # Read markdown
    if args.markdown == '-':
        data = sys.stdin.buffer.read()
        try:
            markdown_text = data.decode('utf-8')
        except UnicodeDecodeError:
            markdown_text = data.decode('utf-8', errors='replace')
    else:
        p = Path(args.markdown)
        markdown_text = p.read_text(encoding='utf-8')

    out_path = Path(args.out).resolve()
    out_path.parent.mkdir(parents=True, exist_ok=True)

    template_path = Path(args.template).resolve()
    block_path = Path(args.template_block).resolve()

    print(f"python-docx version: {getattr(python_docx_module, '__version__', 'unknown')}")
    print("Forbereder dokument…")
    build_doc_from_markdown(markdown_text, template_path, out_path, block_path)
    print("Færdig!")
    print(f"Gemte: {out_path}")


if __name__ == '__main__':
    main()


